import io, os, tempfile, re, cv2
from collections import OrderedDict
from typing import Dict

import streamlit as st
from PIL import Image
import pytesseract
from docx import Document
from docx.shared import Pt

FRAME_INTERVAL = 3            # seconds between OCR frames
st.set_page_config(page_title="Config Extractor", layout="wide")
st.title("🔍 Configuration Extractor → Word")

uploaded = st.file_uploader(
    "Upload an image *or* video of configuration screens:",
    type=["png", "jpg", "jpeg", "bmp", "gif", "tiff", "mp4", "mov", "avi", "mkv"],
)

def ocr_image(img: Image.Image) -> str:
    gray = img.convert("L")
    bw   = gray.point(lambda x: 0 if x < 180 else 255, '1')
    return pytesseract.image_to_string(bw)

def ocr_video(path: str) -> str:
    cap = cv2.VideoCapture(path)
    fps = cap.get(cv2.CAP_PROP_FPS) or 25
    step = int(fps * FRAME_INTERVAL)
    txt = []
    i = 0
    while cap.isOpened():
        ok, frame = cap.read()
        if not ok:
            break
        if i % step == 0:
            pil = Image.fromarray(cv2.cvtColor(frame, cv2.COLOR_BGR2RGB))
            txt.append(ocr_image(pil))
        i += 1
    cap.release()
    return "
".join(txt)

def extract_pairs(raw: str) -> Dict[str, str]:
    pairs = OrderedDict()
    for line in raw.splitlines():
        cleaned = re.sub(r"\s{2,}", "	", line.strip())
        if cleaned.count("	") == 1:
            k, v = map(str.strip, cleaned.split("	"))
            if k and v:
                pairs[k] = v
    return pairs

def to_docx(pairs: Dict[str, str]) -> bytes:
    doc = Document(); doc.add_heading("Configuration Settings", 1)
    table = doc.add_table(rows=1, cols=2)
    hdr1, hdr2 = table.rows[0].cells
    hdr1.text, hdr2.text = "Option", "Value"
    for k, v in pairs.items():
        r = table.add_row().cells
        r[0].text, r[1].text = k, v
    for c in table.rows[0].cells:
        for p in c.paragraphs:
            run = p.runs[0]; run.bold = True; run.font.size = Pt(12)
    bio = io.BytesIO(); doc.save(bio); bio.seek(0); return bio.read()

if uploaded:
    with st.spinner("Processing…"):
        if uploaded.type.startswith("video"):
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".mp4")
            tmp.write(uploaded.read()); tmp.close()
            raw = ocr_video(tmp.name)
            os.unlink(tmp.name)
        else:
            img = Image.open(uploaded).convert("RGB")
            raw = ocr_image(img)
        pairs = extract_pairs(raw)
        if not pairs:
            st.warning("No key/value pairs found. Try a sharper crop.")
        else:
            st.success(f"Found {len(pairs)} settings ✔︎")
            st.json(pairs)
            st.download_button(
                "📥 Download Word Report",
                data=to_docx(pairs),
                file_name="config_report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
```python
